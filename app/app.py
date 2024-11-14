from flask import Flask, render_template, request, send_file
from werkzeug.utils import secure_filename
import os
from docx import Document

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
PROCESSED_FOLDER = 'processed'

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(PROCESSED_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload():
    if 'file' not in request.files:
        return "No file part"
    
    file = request.files['file']
    if file.filename == '':
        return "No selected file"
    
    if file:
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        # Process the Word document
        processed_path = process_word_file(file_path)
        return send_file(processed_path, as_attachment=True)

def process_word_file(file_path):
    # Open and process the Word document
    document = Document(file_path)
    new_doc = Document()
    for paragraph in document.paragraphs:
        new_doc.add_paragraph(paragraph.text.upper())  # Example: Convert text to uppercase

    processed_path = os.path.join(PROCESSED_FOLDER, 'processed_' + os.path.basename(file_path))
    new_doc.save(processed_path)
    return processed_path

if __name__ == '__main__':
    app.run(debug=True)
